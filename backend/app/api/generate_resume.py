"""
POST /generate-resume

Generates a CV (PDF or DOCX) from the candidate's cached resume data.
Data source priority: PostgreSQL sessionactivity snapshots.

PDF  — Playwright renders the generated HTML template to a PDF that mirrors
       the ResumePreview component layout and styling.
DOCX — python-docx builds a professional ATS-friendly document with matching
       sections and structure.
"""
from __future__ import annotations

import html as html_mod
import json
import logging
import re
import uuid
from typing import Annotated, Any

from fastapi import APIRouter, Header, HTTPException, Query
from fastapi.responses import Response
from pydantic import AliasChoices, BaseModel, ConfigDict, Field

from app.core.candidate_cache import _resolve_candidate_id, _cache_key

router = APIRouter(tags=["resume-export"])
logger = logging.getLogger(__name__)


# ── Pydantic request model ────────────────────────────────────────────────────

class GenerateResumeRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    format: str = "pdf"
    career_goal: str = ""
    region: str = ""
    completed_skills: list[str] = Field(default_factory=list)
    completed_ai_skills: list[str] = Field(
        default_factory=list,
        validation_alias=AliasChoices("completed_ai_skills", "completedAiSkills", "completedAISkills"),
    )
    completed_competencies: list[str] = Field(default_factory=list)
    completed_certifications: list[str] = Field(default_factory=list)


# ── Data fetching — PostgreSQL only ──────────────────────────────────────────

def _read_frontend_resume_cache(session_token: str) -> dict | None:
    """Stub — Redis removed. Always returns None."""
    return None


def _read_candidate_cache(candidate_id: uuid.UUID, session_token: str) -> dict | None:
    """Stub — Redis removed. Always returns None."""
    return None


async def _read_session_activity_pg(candidate_id: uuid.UUID) -> dict | None:
    """Fallback: read the most recent sessionactivity snapshot from PostgreSQL."""
    try:
        from app.core.async_db import get_async_pool
        pool = get_async_pool()
        async with pool.connection() as conn:
            cur = await conn.execute(
                """
                SELECT full_cache_snapshot FROM sessionactivity
                WHERE candidate_id = %s
                ORDER BY logged_out_at DESC NULLS LAST
                LIMIT 1
                """,
                (str(candidate_id),),
            )
            row = await cur.fetchone()
            if row and row[0]:
                data = row[0]
                if isinstance(data, str):
                    try:
                        return json.loads(data)
                    except json.JSONDecodeError:
                        return {"rawCvText": data}
                return data
    except Exception as exc:
        logger.warning("[generate-resume] PG sessionactivity read failed: %s", exc)
    return None


async def _read_cv_upload_pg(candidate_id: uuid.UUID) -> dict | None:
    """Read the most recent CV upload from user_cv_upload (primary data source)."""
    try:
        from app.core.async_db import get_async_pool
        pool = get_async_pool()
        async with pool.connection() as conn:
            cur = await conn.execute(
                """
                SELECT contact_info, work_experience, education, sections,
                       core_skills, tools_and_technologies, certifications,
                       years_of_experience, location, raw_text, future_roles_data,
                       profile_summary, parsed_output
                FROM user_cv_upload
                WHERE candidate_id = %s
                ORDER BY created_at DESC
                LIMIT 1
                """,
                (str(candidate_id),),
            )
            row = await cur.fetchone()
            if not row:
                return None
            cols = [
                "contact_info", "work_experience", "education", "sections",
                "core_skills", "tools_and_technologies", "certifications",
                "years_of_experience", "location", "raw_text", "future_roles_data",
                "profile_summary", "parsed_output",
            ]
            data: dict[str, Any] = {}
            for i, col in enumerate(cols):
                val = row[i]
                if isinstance(val, str) and col not in ("education", "location", "raw_text"):
                    try:
                        val = json.loads(val)
                    except (json.JSONDecodeError, TypeError):
                        pass
                data[col] = val
            data["_source"] = "cv_upload"
            return data
    except Exception as exc:
        logger.warning("[generate-resume] PG cv_upload read failed: %s", exc)
    return None


# ── Resume data structuring ────────────────────────────────────────────────────

def _parse_experience(text: Any) -> list[dict]:
    """Parse work_experience text → list of structured experience dicts."""
    if isinstance(text, list):
        return text
    if not text or not isinstance(text, str):
        return []
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    experiences: list[dict] = []
    current: dict | None = None
    for line in lines:
        is_bullet = bool(re.match(r"^[•\-–—*►▸]\s*", line)) or bool(re.match(r"^\d+[.)]\s", line))
        if is_bullet and current:
            clean = re.sub(r"^[•\-–—*►▸]\s*", "", line)
            clean = re.sub(r"^\d+[.)]\s", "", clean).strip()
            if clean:
                current["bullets"].append(clean)
            continue
        has_date = bool(re.search(r"\b(19|20)\d{2}\b", line)) or bool(re.search(r"\b(present|current)\b", line, re.I))
        has_sep = bool(re.search(r"\s[—–\-|]\s", line)) or "," in line
        is_short = len(line) < 120
        if (has_date or has_sep) and is_short and not is_bullet:
            m = re.match(r"^(.+?)\s*[—–\-|]\s*(.+?)(?:\s{2,}|\s*,\s*)(.+)$", line)
            if m:
                title, company, dates = m.group(1).strip(), m.group(2).strip(), m.group(3).strip()
            else:
                parts = re.split(r"\s{2,}", line)
                title = parts[0] if parts else line
                dates = parts[-1] if len(parts) > 1 else ""
                company = parts[1] if len(parts) > 2 else ""
            if current:
                experiences.append(current)
            current = {"title": title, "company": company, "location": "", "dates": dates, "bullets": []}
        elif current:
            current["bullets"].append(line)
        else:
            current = {"title": line, "company": "", "location": "", "dates": "", "bullets": []}
    if current:
        experiences.append(current)
    return experiences


def _parse_education(text: Any) -> list[dict]:
    """Parse education text → list of structured education dicts."""
    if isinstance(text, list):
        return text
    if not text or not isinstance(text, str):
        return []
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    education: list[dict] = []
    for line in lines:
        if re.match(r"^[•\-–—*►▸]\s*", line):
            continue
        m = re.match(r"^(.+?)\s*[—–\-|]\s*(.+?)(?:\s{2,}|\s*,\s*)(.+)$", line)
        if m:
            degree, institution, dates = m.group(1).strip(), m.group(2).strip(), m.group(3).strip()
        else:
            parts = re.split(r"\s{2,}", line)
            degree = parts[0] if parts else line
            dates = parts[-1] if len(parts) > 1 else ""
            institution = parts[1] if len(parts) > 2 else ""
        if degree:
            education.append({"degree": degree, "institution": institution, "dates": dates})
    return education


def _parse_projects(text: Any) -> list[str]:
    if isinstance(text, list):
        return [str(p) for p in text if p]
    if not text or not isinstance(text, str):
        return []
    return [re.sub(r"^[•\-–—*►▸]\s*", "", l).strip()
            for l in text.split("\n") if l.strip()]


def _extract_contact_from_text(raw: str) -> dict:
    contact: dict[str, str] = {}
    if not raw:
        return contact
    lines = [l.strip() for l in raw.split("\n") if l.strip()]
    for line in lines[:5]:
        if 2 < len(line) < 60 and not re.search(r"[@.\d]{4,}", line) and not re.match(r"^(contact|phone|email|address|mobile)", line, re.I):
            contact["name"] = line
            break
    em = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", raw)
    if em:
        contact["email"] = em.group(0)
    ph = re.search(r"(?:\+?\d{1,3}[\s.-]?)?\(?\d{2,4}\)?[\s.-]?\d{3,4}[\s.-]?\d{3,4}", raw)
    if ph:
        contact["phone"] = ph.group(0).strip()
    li = re.search(r"(?:linkedin\.com/in/[\w-]+|linkedin:\s*[\w/.-]+)", raw, re.I)
    if li:
        contact["linkedin"] = li.group(0)
    return contact


def _extract_summary(raw: str) -> str:
    if not raw:
        return ""
    summary_re = [
        re.compile(r"(?:personal\s+)?(?:profile|summary|statement|objective)", re.I),
        re.compile(r"professional\s+summary", re.I),
        re.compile(r"about\s+me", re.I),
        re.compile(r"career\s+(?:summary|objective|profile)", re.I),
    ]
    section_re = [
        re.compile(r"\b(?:work|professional)\s+(?:experience|history)\b", re.I),
        re.compile(r"\beducation\b", re.I),
        re.compile(r"\bcertification", re.I),
        re.compile(r"\bskills?\b", re.I),
    ]
    lines = raw.split("\n")
    capturing = False
    result: list[str] = []
    for line in lines:
        line = line.strip()
        if not line:
            if capturing and result:
                break
            continue
        if not capturing:
            if any(r.search(line) for r in summary_re):
                capturing = True
                continue
        else:
            if any(r.search(line) for r in section_re):
                break
            result.append(line)
    return " ".join(result).strip()


def build_structured_resume(
    cached: dict,
    career_goal: str = "",
    region: str = "",
    completed_skills: list[str] | None = None,
    completed_ai_skills: list[str] | None = None,
    completed_competencies: list[str] | None = None,
    completed_certifications: list[str] | None = None,
) -> dict:
    """
    Python equivalent of buildStructuredResume from parseCvSections.js.
    Handles three data formats:
      1. user_cv_upload (primary — structured JSONB columns)
      2. Frontend cache (legacy — Redis-era format)
      3. sessionactivity snapshot (fallback — full_cache_snapshot JSONB)
    """

    # ── user_cv_upload data path ──────────────────────────────────────────────
    if cached.get("_source") == "cv_upload":
        return _build_from_cv_upload(
            cached, career_goal, region,
            completed_skills, completed_ai_skills, completed_competencies, completed_certifications,
        )

    # ── Legacy / sessionactivity data path (unchanged) ────────────────────────
    raw_text = cached.get("rawCvText") or cached.get("raw_cv_text") or ""
    contact_from_text = _extract_contact_from_text(raw_text)

    cd = cached.get("CandidateDetails") or {}

    summary = cached.get("summary") or _extract_summary(raw_text)

    core_skills_raw = cached.get("coreSkills") or cached.get("core_skills") or []
    skill_categories = cached.get("skillCategories") or cached.get("core_skills") or None
    if isinstance(core_skills_raw, dict):
        core_skills = list(core_skills_raw.keys())
        skill_categories = core_skills_raw
    elif isinstance(core_skills_raw, list):
        core_skills = core_skills_raw
    else:
        core_skills = []

    tools_raw = cached.get("tools") or cached.get("tools_and_technologies") or []
    if isinstance(tools_raw, dict):
        tools = list(tools_raw.keys())
    elif isinstance(tools_raw, list):
        tools = tools_raw
    else:
        tools = []

    experience = _parse_experience(cached.get("experience") or cached.get("work_experience") or "")
    education = _parse_education(cached.get("education") or "")
    projects = _parse_projects(cached.get("projects") or "")

    certs_raw = cached.get("certifications") or []
    certs = [str(c) for c in certs_raw if c] if isinstance(certs_raw, list) else []

    all_skills = list(dict.fromkeys(core_skills + (completed_skills or [])))
    all_ai_skills = list(dict.fromkeys(completed_ai_skills or []))
    all_certs = list(dict.fromkeys(certs + (completed_certifications or [])))
    all_comps = list(completed_competencies or [])

    return {
        "contact": {
            "name": cd.get("Name") or contact_from_text.get("name") or cached.get("contact", {}).get("name") or "",
            "headline": career_goal or cached.get("bestFitIndustry") or cached.get("best_fit_industry") or "",
            "email": cd.get("Email") or contact_from_text.get("email") or cached.get("contact", {}).get("email") or "",
            "phone": cd.get("Phone") or contact_from_text.get("phone") or cached.get("contact", {}).get("phone") or "",
            "location": region or cached.get("contact", {}).get("location") or "",
            "linkedin": contact_from_text.get("linkedin") or "",
            "github": "",
            "website": "",
        },
        "summary": summary,
        "coreSkills": all_skills,
        "aiSkills": all_ai_skills,
        "skillCategories": skill_categories if isinstance(skill_categories, dict) and skill_categories else None,
        "tools": tools,
        "toolCategories": None,
        "competencies": all_comps,
        "certifications": all_certs,
        "experience": experience,
        "education": education,
        "projects": projects,
    }


def _build_from_cv_upload(
    cached: dict,
    career_goal: str,
    region: str,
    completed_skills: list[str] | None,
    completed_ai_skills: list[str] | None,
    completed_competencies: list[str] | None,
    completed_certifications: list[str] | None,
) -> dict:
    """Build structured resume dict from user_cv_upload row data."""
    contact_info = cached.get("contact_info") or {}
    sections = cached.get("sections") or {}

    # Contact
    contact = {
        "name": contact_info.get("name", ""),
        "headline": career_goal or "",
        "email": contact_info.get("email", ""),
        "phone": contact_info.get("phone", ""),
        "location": region or cached.get("location", "") or "",
        "linkedin": contact_info.get("linkedin", ""),
        "github": contact_info.get("github", ""),
        "website": contact_info.get("website", ""),
    }

    # Summary - prefer profile_summary from database
    summary = cached.get("profile_summary") or sections.get("summary", "")

    # Work experience — use structured entries from JSONB
    work_exp = cached.get("work_experience") or {}
    entries = work_exp.get("entries", []) if isinstance(work_exp, dict) else []
    experience: list[dict] = []
    for entry in entries:
        date_range = entry.get("date_range") or {}
        dates_str = date_range.get("raw", "")
        if not dates_str and date_range.get("start_date"):
            end = "Present" if date_range.get("is_current") else (date_range.get("end_date") or "")
            dates_str = f"{date_range['start_date']} - {end}"
        desc = entry.get("description") or ""
        bullets = [b.strip() for b in desc.split("\n") if b.strip()]
        experience.append({
            "title": entry.get("title", ""),
            "company": entry.get("company", ""),
            "location": entry.get("location", ""),
            "dates": dates_str,
            "bullets": bullets,
        })
    if not experience:
        experience = _parse_experience(work_exp.get("formatted_text", "") if isinstance(work_exp, dict) else "")

    # Education
    education = _parse_education(cached.get("education") or "")

    # Core skills — invert {skill: category} → {category: [skills]}
    core_skills_dict = cached.get("core_skills") or {}
    core_skills: list[str] = []
    skill_categories: dict[str, list[str]] = {}
    if isinstance(core_skills_dict, dict) and core_skills_dict:
        core_skills = list(core_skills_dict.keys())
        for skill, cat in core_skills_dict.items():
            skill_categories.setdefault(cat, []).append(skill)

    # Tools — invert {tool: category} → {category: [tools]}
    tools_dict = cached.get("tools_and_technologies") or {}
    tools: list[str] = []
    tool_categories: dict[str, list[str]] = {}
    if isinstance(tools_dict, dict) and tools_dict:
        tools = list(tools_dict.keys())
        for tool, cat in tools_dict.items():
            tool_categories.setdefault(cat, []).append(tool)

    # Certifications
    certs_raw = cached.get("certifications") or []
    certs = [str(c) for c in certs_raw if c] if isinstance(certs_raw, list) else []

    # Projects
    projects = _parse_projects(sections.get("projects", ""))

    # Merge completed items from frontend checkboxes
    all_skills = list(dict.fromkeys(core_skills + (completed_skills or [])))
    all_ai_skills = list(dict.fromkeys(completed_ai_skills or []))
    all_certs = list(dict.fromkeys(certs + (completed_certifications or [])))
    all_comps = list(completed_competencies or [])

    # Pass through raw database data for frontend to use
    db_data = {
        "profile_summary": cached.get("profile_summary"),
        "work_experience": cached.get("work_experience"),
        "education": cached.get("education"),
        "raw_text": cached.get("raw_text"),
        "contact_info": contact_info,
        "sections": sections,
        "parsed_output": cached.get("parsed_output"),
    }

    return {
        "contact": contact,
        "summary": summary,
        "coreSkills": all_skills,
        "aiSkills": all_ai_skills,
        "skillCategories": skill_categories if skill_categories else None,
        "tools": tools,
        "toolCategories": tool_categories if tool_categories else None,
        "competencies": all_comps,
        "certifications": all_certs,
        "experience": experience,
        "education": education,
        "projects": projects,
        "dbData": db_data,  # Add raw database data for frontend
    }


# ── HTML builder (PDF) — premium ATS-friendly executive template ─────────────

def _e(v: Any) -> str:
    return html_mod.escape(str(v or ""))


def _skill_badge(text: str, tint: str = "indigo") -> str:
    return f'<span class="badge badge-{tint}">{_e(text)}</span>'


def _bullet_item(text: str, dot_class: str = "") -> str:
    dc = f"bi-dot {dot_class}" if dot_class else "bi-dot"
    return f'<div class="bi"><span class="{dc}"></span><span class="bi-txt">{_e(text)}</span></div>'


def _section_heading(title: str, line_class: str = "sec-line-indigo") -> str:
    return f'<h2>{_e(title)}</h2><div class="sec-line {line_class}"></div>'


PDF_CSS = """
@page{size:A4;margin:0;}
*{box-sizing:border-box;margin:0;padding:0;print-color-adjust:exact;-webkit-print-color-adjust:exact;}
body{font-family:Calibri,'Segoe UI',Arial,Helvetica,sans-serif;color:#1e293b;font-size:10.5px;line-height:1.5;}

/* Top accent bar */
.accent-bar{height:3px;background:linear-gradient(to right,#6366f1,#a855f7,#ec4899);}

/* Header — matches ResumePreview dark gradient */
.hdr{background:linear-gradient(135deg,#0f172a 0%,#1e293b 50%,#312e81 100%);padding:26px 30px 22px;color:#fff;position:relative;}
.hdr h1{font-size:26px;font-weight:700;letter-spacing:-0.01em;line-height:1.15;}
.hdr .headline{font-size:13px;font-weight:500;color:rgba(165,180,252,0.9);margin-top:3px;}
.hdr .contact-row{display:flex;flex-wrap:wrap;align-items:center;gap:4px;margin-top:10px;font-size:10px;color:#cbd5e1;font-weight:400;}
.hdr .contact-pill{background:rgba(255,255,255,0.07);padding:3px 10px;border-radius:20px;display:inline-flex;align-items:center;gap:4px;}
.hdr .sep{color:rgba(129,140,248,0.4);margin:0 2px;}

/* Two-column body */
.body-grid{display:grid;grid-template-columns:30% 1fr;min-height:800px;}
.sidebar{background:linear-gradient(to bottom,#f8fafc,#ffffff,rgba(248,250,252,0.8));border-right:1px solid rgba(226,232,240,0.7);padding:22px 18px;}
.main{padding:22px 24px;}

/* Section */
.sec{margin-bottom:16px;}
.sec h2{font-size:10px;font-weight:700;letter-spacing:0.12em;text-transform:uppercase;color:#1e293b;margin-bottom:6px;}
.sec-line{height:1px;margin-bottom:8px;opacity:0.2;}
.sec-line-indigo{background:linear-gradient(to right,#6366f1,#3b82f6);}
.sec-line-cyan{background:linear-gradient(to right,#06b6d4,#14b8a6);}
.sec-line-amber{background:linear-gradient(to right,#f59e0b,#f97316);}
.sec-line-violet{background:linear-gradient(to right,#8b5cf6,#a855f7);}
.sec-line-slate{background:linear-gradient(to right,#475569,#64748b);}
.sec-line-blue{background:linear-gradient(to right,#3b82f6,#06b6d4);}
.sec-line-rose{background:linear-gradient(to right,#f43f5e,#ec4899);}
.sec-line-emerald{background:linear-gradient(to right,#10b981,#14b8a6);}

/* Skill badges — color-coded to match ResumePreview */
.badge{display:inline-block;padding:2px 8px;border-radius:5px;font-size:9px;font-weight:600;margin:0 4px 5px 0;border:1px solid;}
.badge-indigo{background:rgba(238,242,255,0.8);border-color:rgba(199,210,254,0.6);color:#4338ca;}
.badge-cyan{background:rgba(236,254,255,0.8);border-color:rgba(165,243,252,0.6);color:#0e7490;}
.badge-amber{background:rgba(255,251,235,0.8);border-color:rgba(253,230,138,0.6);color:#92400e;}
.badge-violet{background:rgba(245,243,255,0.8);border-color:rgba(221,214,254,0.6);color:#6d28d9;}
.cat-label{font-size:9px;font-weight:700;text-transform:uppercase;letter-spacing:0.06em;margin:8px 0 3px 0;}
.cat-label:first-child{margin-top:0;}
.cat-label-indigo{color:rgba(99,102,241,0.7);}
.cat-label-cyan{color:rgba(6,182,212,0.7);}

/* Bullet items — gradient dot matching ResumePreview */
.bi{display:flex;align-items:flex-start;gap:8px;margin-bottom:4px;}
.bi-dot{width:5px;height:5px;border-radius:50%;background:linear-gradient(135deg,#818cf8,#a855f7);margin-top:5px;flex-shrink:0;}
.bi-dot-violet{background:linear-gradient(135deg,#8b5cf6,#a855f7);}
.bi-txt{font-size:10px;line-height:1.6;color:#475569;font-weight:500;}

/* Experience — Timeline layout matching ResumePreview */
.exp-timeline{padding-left:18px;border-left:2px solid #e0e7ff;margin-left:4px;}
.exp{position:relative;margin-bottom:10px;break-inside:avoid;}
.exp-dot{position:absolute;left:-24px;top:10px;width:10px;height:10px;border-radius:50%;background:linear-gradient(135deg,#6366f1,#a855f7);border:2px solid white;}
.exp-card{background:rgba(248,250,252,0.6);border-radius:8px;padding:10px 12px;border:1px solid rgba(241,245,249,0.8);}
.exp-top{display:flex;flex-wrap:wrap;align-items:flex-start;justify-content:space-between;gap:6px;}
.exp-title{font-size:11px;font-weight:700;color:#0f172a;line-height:1.3;flex:1;min-width:0;}
.exp-dates{font-size:9px;font-weight:600;color:#64748b;background:white;padding:2px 8px;border-radius:20px;border:1px solid rgba(226,232,240,0.8);white-space:nowrap;flex-shrink:0;}
.exp-company{font-size:9.5px;font-weight:600;color:rgba(79,70,229,0.7);margin-top:2px;margin-bottom:4px;}
.exp-bullets{margin-top:6px;}

/* Education — card style matching ResumePreview */
.edu-card{background:rgba(239,246,255,0.3);border-radius:8px;padding:8px 12px;border:1px solid rgba(219,234,254,0.5);margin-bottom:6px;break-inside:avoid;}
.edu-top{display:flex;flex-wrap:wrap;align-items:flex-start;justify-content:space-between;gap:6px;}
.edu-degree{font-size:11px;font-weight:700;color:#1e293b;}
.edu-inst{font-size:10px;font-weight:600;color:rgba(37,99,235,0.7);margin-top:2px;}

/* Certification items */
.cert-item{display:flex;align-items:flex-start;gap:6px;margin-bottom:6px;}
.cert-icon{color:#f59e0b;font-size:10px;margin-top:1px;flex-shrink:0;}
.cert-txt{font-size:10px;line-height:1.55;color:#334155;font-weight:500;}

/* Summary */
.summary-txt{font-size:10.5px;line-height:1.7;color:#475569;font-weight:500;}
"""


def _build_resume_html(resume: dict) -> str:
    contact = resume.get("contact", {})
    name = contact.get("name") or "Your Name"
    headline = contact.get("headline") or ""
    email = contact.get("email") or ""
    phone = contact.get("phone") or ""
    location = contact.get("location") or ""
    linkedin = contact.get("linkedin") or ""
    github = contact.get("github") or ""
    website = contact.get("website") or ""

    # Contact row — pill-shaped items separated by pipe (matches ResumePreview)
    contact_parts: list[str] = []
    if email:
        contact_parts.append(f'<span class="contact-pill">{_e(email)}</span>')
    if phone:
        contact_parts.append(f'<span class="contact-pill">{_e(phone)}</span>')
    if location:
        contact_parts.append(f'<span class="contact-pill">{_e(location)}</span>')
    if linkedin:
        contact_parts.append(f'<span class="contact-pill">{_e(linkedin)}</span>')
    if github:
        contact_parts.append(f'<span class="contact-pill">{_e(github)}</span>')
    if website:
        contact_parts.append(f'<span class="contact-pill">{_e(website)}</span>')
    contact_html = '<span class="sep">|</span>'.join(contact_parts)

    # ── Sidebar ──
    sidebar = ""

    # Core Skills (grouped by category — indigo tint)
    core_skills = resume.get("coreSkills") or []
    skill_categories = resume.get("skillCategories")
    if core_skills:
        sidebar += f'<div class="sec">{_section_heading("Core Skills", "sec-line-indigo")}'
        if skill_categories and isinstance(skill_categories, dict):
            for cat, skills in skill_categories.items():
                sidebar += f'<p class="cat-label cat-label-indigo">{_e(cat)}</p>'
                sk_list = skills if isinstance(skills, list) else [skills]
                sidebar += "".join(_skill_badge(s, "indigo") for s in sk_list if s)
        else:
            sidebar += "".join(_skill_badge(s, "indigo") for s in core_skills if s)
        sidebar += "</div>"

    # Tools (grouped by category — cyan tint)
    tools = resume.get("tools") or []
    tool_categories = resume.get("toolCategories")
    if tools:
        sidebar += f'<div class="sec">{_section_heading("Tools & Technologies", "sec-line-cyan")}'
        if tool_categories and isinstance(tool_categories, dict):
            for cat, items in tool_categories.items():
                sidebar += f'<p class="cat-label cat-label-cyan">{_e(cat)}</p>'
                sidebar += "".join(_skill_badge(t, "cyan") for t in items if t)
        else:
            sidebar += "".join(_skill_badge(t, "cyan") for t in tools if t)
        sidebar += "</div>"

    # Certifications (amber tint)
    certs = resume.get("certifications") or []
    if certs:
        sidebar += f'<div class="sec">{_section_heading("Certifications", "sec-line-amber")}'
        for c in certs:
            if c:
                sidebar += f'<div class="cert-item"><span class="cert-icon">&#9733;</span><span class="cert-txt">{_e(c)}</span></div>'
        sidebar += "</div>"

    # Competencies (violet tint)
    comps = resume.get("competencies") or []
    if comps:
        sidebar += f'<div class="sec">{_section_heading("Competencies", "sec-line-violet")}'
        sidebar += "".join(_bullet_item(c, "bi-dot-violet") for c in comps if c)
        sidebar += "</div>"

    # AI Skills (teal tint)
    ai_skills = resume.get("aiSkills") or []
    if ai_skills:
        sidebar += f'<div class="sec">{_section_heading("AI Skills", "sec-line-cyan")}'
        sidebar += "".join(_skill_badge(s, "cyan") for s in ai_skills if s)
        sidebar += "</div>"

    # ── Main content ──
    main = ""

    # Summary
    summary = resume.get("summary") or ""
    if summary:
        main += f'<div class="sec">{_section_heading("Professional Summary", "sec-line-slate")}<p class="summary-txt">{_e(summary)}</p></div>'

    # Experience — Timeline layout matching ResumePreview
    experience = resume.get("experience") or []
    if experience:
        main += f'<div class="sec">{_section_heading("Professional Experience", "sec-line-indigo")}'
        main += '<div class="exp-timeline">'
        for exp in experience:
            if isinstance(exp, dict):
                title = _e(exp.get("title") or "")
                company = exp.get("company") or ""
                loc = exp.get("location") or ""
                dates = _e(exp.get("dates") or "")
                company_parts = [_e(company)] if company else []
                if loc:
                    company_parts.append(_e(loc))
                company_line = " • ".join(company_parts)
                main += '<div class="exp"><div class="exp-dot"></div><div class="exp-card">'
                main += f'<div class="exp-top"><span class="exp-title">{title}</span>'
                if dates:
                    main += f'<span class="exp-dates">{dates}</span>'
                main += "</div>"
                if company_line:
                    main += f'<p class="exp-company">{company_line}</p>'
                if exp.get("bullets"):
                    main += '<div class="exp-bullets">'
                    for b in exp["bullets"]:
                        main += _bullet_item(b)
                    main += '</div>'
                main += "</div></div>"
            else:
                main += _bullet_item(str(exp))
        main += "</div></div>"

    # Education — card style matching ResumePreview
    education = resume.get("education") or []
    if education:
        main += f'<div class="sec">{_section_heading("Education", "sec-line-blue")}'
        for edu in education:
            if isinstance(edu, dict):
                degree = _e(edu.get("degree") or "")
                institution = edu.get("institution") or ""
                dates = _e(edu.get("dates") or "")
                main += '<div class="edu-card"><div class="edu-top"><div>'
                main += f'<p class="edu-degree">{degree}</p>'
                if institution:
                    main += f'<p class="edu-inst">{_e(institution)}</p>'
                main += "</div>"
                if dates:
                    main += f'<span class="exp-dates">{dates}</span>'
                main += "</div></div>"
            else:
                main += _bullet_item(str(edu))
        main += "</div>"

    # Projects
    projects = resume.get("projects") or []
    if projects:
        main += f'<div class="sec">{_section_heading("Projects", "sec-line-rose")}'
        main += "".join(_bullet_item(p) for p in projects if p)
        main += "</div>"

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>{_e(name)} Resume</title>
<style>{PDF_CSS}</style>
</head>
<body>
<div class="accent-bar"></div>
<div class="hdr">
  <h1>{_e(name)}</h1>
  {"<p class='headline'>" + _e(headline) + "</p>" if headline else ""}
  <div class="contact-row">{contact_html}</div>
</div>
<div class="body-grid">
  <div class="sidebar">{sidebar}</div>
  <div class="main">{main}</div>
</div>
</body>
</html>"""


# ── DOCX builder (python-docx) — premium two-column layout ───────────────────

def _build_resume_docx(resume: dict) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import io

    NAVY = RGBColor(0x1E, 0x29, 0x3B)
    SLATE_700 = RGBColor(0x33, 0x41, 0x55)
    SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
    DARK = RGBColor(0x1E, 0x29, 0x3B)

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = DARK

    # Reduce default margins for more space
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    contact = resume.get("contact", {})
    name = contact.get("name") or "Your Name"
    headline = contact.get("headline") or ""

    # ── Header ──
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2)

    if headline:
        p = doc.add_paragraph()
        run = p.add_run(headline)
        run.font.size = Pt(11)
        run.font.color.rgb = SLATE_500
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(2)

    contact_parts = [v for v in [
        contact.get("email"), contact.get("phone"), contact.get("location"),
        contact.get("linkedin"), contact.get("github"), contact.get("website"),
    ] if v]
    if contact_parts:
        p = doc.add_paragraph()
        run = p.add_run("  |  ".join(contact_parts))
        run.font.size = Pt(8)
        run.font.color.rgb = SLATE_500
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(6)

    # Horizontal rule below header
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single", qn("w:sz"): "4",
        qn("w:space"): "1", qn("w:color"): "CBD5E1",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Two-column table ──
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(5.0)

    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(tblPr.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "none", qn("w:sz"): "0", qn("w:space"): "0", qn("w:color"): "auto",
        }))
    existing_borders = tblPr.find(qn("w:tblBorders"))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    # Sidebar shading
    left_cell = table.cell(0, 0)
    tc = left_cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn("w:shd"), {
        qn("w:fill"): "F8FAFC", qn("w:val"): "clear",
    })
    tcPr.append(shading)

    right_cell = table.cell(0, 1)

    # Helper functions for cell content
    def add_cell_heading(cell, text):
        p = cell.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = SLATE_700
        run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        btm = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "4",
            qn("w:space"): "1", qn("w:color"): "CBD5E1",
        })
        pBdr.append(btm)
        pPr.append(pBdr)

    def add_cell_text(cell, text, bold=False, color=DARK, size=9):
        p = cell.add_paragraph()
        run = p.add_run(str(text or ""))
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(1)
        return p

    def add_cell_bullet(cell, text, size=9):
        p = cell.add_paragraph()
        run = p.add_run(f"•  {text}")
        run.font.size = Pt(size)
        run.font.color.rgb = DARK
        run.font.name = "Calibri"
        p.paragraph_format.left_indent = Pt(10)
        p.paragraph_format.space_after = Pt(1)

    # Remove default empty paragraph from cells
    for cell in (left_cell, right_cell):
        if cell.paragraphs and not cell.paragraphs[0].text:
            p_elem = cell.paragraphs[0]._p
            p_elem.getparent().remove(p_elem)

    # ── LEFT COLUMN (sidebar) ──
    # Skills
    skill_categories = resume.get("skillCategories")
    core_skills = resume.get("coreSkills") or []
    if skill_categories and isinstance(skill_categories, dict):
        add_cell_heading(left_cell, "Skills")
        for cat, skills in skill_categories.items():
            add_cell_text(left_cell, cat, bold=True, color=SLATE_700, size=8)
            sk = skills if isinstance(skills, list) else [skills]
            add_cell_text(left_cell, "  •  ".join(str(s) for s in sk if s), size=8)
    elif core_skills:
        add_cell_heading(left_cell, "Skills")
        add_cell_text(left_cell, "  •  ".join(str(s) for s in core_skills if s), size=8)

    # Tools
    tools = resume.get("tools") or []
    tool_categories = resume.get("toolCategories")
    if tool_categories and isinstance(tool_categories, dict):
        add_cell_heading(left_cell, "Tools & Technologies")
        for cat, items in tool_categories.items():
            add_cell_text(left_cell, cat, bold=True, color=SLATE_700, size=8)
            add_cell_text(left_cell, "  •  ".join(str(t) for t in items if t), size=8)
    elif tools:
        add_cell_heading(left_cell, "Tools & Technologies")
        add_cell_text(left_cell, "  •  ".join(str(t) for t in tools if t), size=8)

    # Certifications
    certs = resume.get("certifications") or []
    if certs:
        add_cell_heading(left_cell, "Certifications")
        for c in certs:
            if c:
                add_cell_bullet(left_cell, str(c), size=8)

    # Competencies
    comps = resume.get("competencies") or []
    if comps:
        add_cell_heading(left_cell, "Competencies")
        for c in comps:
            if c:
                add_cell_bullet(left_cell, str(c), size=8)

    # AI Skills
    ai_skills = resume.get("aiSkills") or []
    if ai_skills:
        add_cell_heading(left_cell, "AI Skills")
        add_cell_text(left_cell, "  •  ".join(str(s) for s in ai_skills if s), size=8)

    # ── RIGHT COLUMN (main content) ──
    # Summary
    summary = resume.get("summary") or ""
    if summary:
        add_cell_heading(right_cell, "Professional Summary")
        add_cell_text(right_cell, summary, size=10)

    # Experience
    experience = resume.get("experience") or []
    if experience:
        add_cell_heading(right_cell, "Professional Experience")
        for exp in experience:
            if isinstance(exp, dict):
                title = exp.get("title") or ""
                company = exp.get("company") or ""
                dates = exp.get("dates") or ""
                loc = exp.get("location") or ""
                title_line = title
                if dates:
                    title_line += f"    ({dates})"
                add_cell_text(right_cell, title_line, bold=True, size=10)
                company_parts = [company] if company else []
                if loc:
                    company_parts.append(loc)
                if company_parts:
                    add_cell_text(right_cell, ", ".join(company_parts), color=SLATE_500, size=9)
                for b in (exp.get("bullets") or []):
                    add_cell_bullet(right_cell, str(b))
                # Small spacer between entries
                p = right_cell.add_paragraph()
                p.paragraph_format.space_after = Pt(2)

    # Education
    education = resume.get("education") or []
    if education:
        add_cell_heading(right_cell, "Education")
        for edu in education:
            if isinstance(edu, dict):
                line = edu.get("degree") or ""
                if edu.get("institution"):
                    line += f" — {edu['institution']}"
                if edu.get("dates"):
                    line += f"    ({edu['dates']})"
                add_cell_text(right_cell, line, size=10)

    # Projects
    projects = resume.get("projects") or []
    if projects:
        add_cell_heading(right_cell, "Projects")
        for p_item in projects:
            if p_item:
                add_cell_bullet(right_cell, str(p_item))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF renderer ────────────────────────────────────────────────────────────

from app.api.pdf_renderer import render_html_to_pdf_async


async def _html_to_pdf(html_content: str) -> bytes:
    return await render_html_to_pdf_async(
        html_content,
        media="print",
        viewport_width=780,
        viewport_height=1200,
        format="A4",
        margin={"top": "0", "right": "0", "bottom": "0", "left": "0"},
    )


# ── Endpoint ───────────────────────────────────────────────────────────────────

@router.post("/generate-resume")
async def generate_resume(
    req: GenerateResumeRequest,
    x_session_id: Annotated[str | None, Header()] = None,
) -> Response:
    """Generate and download a CV (PDF or DOCX) from cached resume data."""
    session_token = x_session_id
    if not session_token:
        raise HTTPException(status_code=400, detail="x-session-id header is required")

    fmt = req.format.lower().strip()
    if fmt not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}. Use 'pdf' or 'docx'.")

    # ── 1. Fetch resume data — cv_upload (primary), sessionactivity (fallback) ──
    cached: dict | None = None
    cid = _resolve_candidate_id(session_token)
    if cid:
        try:
            cached = await _read_cv_upload_pg(cid)
            if cached:
                logger.info("[generate-resume] HIT PG cv_upload — cid=%s", cid)
        except Exception as exc:
            logger.warning("[generate-resume] cv_upload lookup error: %s", exc)

        if not cached:
            try:
                cached = await _read_session_activity_pg(cid)
                if cached:
                    logger.info("[generate-resume] HIT PG sessionactivity — cid=%s", cid)
            except Exception as exc:
                logger.warning("[generate-resume] sessionactivity lookup error: %s", exc)

    if not cached:
        return Response(
            content=json.dumps({"success": False, "detail": "No resume data found in cache for this session. Please upload a CV first."}),
            media_type="application/json",
            status_code=404,
        )

    # ── 2. Structure the resume data ──
    resume = build_structured_resume(
        cached,
        career_goal=req.career_goal,
        region=req.region,
        completed_skills=req.completed_skills or None,
        completed_ai_skills=req.completed_ai_skills or None,
        completed_competencies=req.completed_competencies or None,
        completed_certifications=req.completed_certifications or None,
    )

    safe_name = re.sub(r"[^a-zA-Z0-9_\- ]", "", resume["contact"].get("name") or "Resume").strip().replace(" ", "_") or "Resume"

    # ── 3. Generate the file ──
    try:
        if fmt == "pdf":
            html_doc = _build_resume_html(resume)
            file_bytes = await _html_to_pdf(html_doc)
            media_type = "application/pdf"
            filename = f"{safe_name}_Resume.pdf"
        else:
            file_bytes = _build_resume_docx(resume)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{safe_name}_Resume.docx"
    except Exception as exc:
        logger.exception("[generate-resume] generation failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"{type(exc).__name__}: {exc}") from exc

    return Response(
        content=file_bytes,
        media_type=media_type,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Content-Type": media_type,
        },
    )


@router.get("/generate-resume-preview")
async def generate_resume_preview(
    x_session_id: Annotated[str | None, Header()] = None,
    career_goal: str = "",
    region: str = "",
) -> dict:
    """Return structured resume data as JSON for the frontend preview modal."""
    session_token = x_session_id
    if not session_token:
        raise HTTPException(status_code=400, detail="x-session-id header is required")

    cached: dict | None = None
    cid = _resolve_candidate_id(session_token)
    if cid:
        cached = await _read_cv_upload_pg(cid)
        if cached:
            logger.info("[generate-resume-preview] HIT PG cv_upload — cid=%s", cid)

        if not cached:
            cached = await _read_session_activity_pg(cid)
            if cached:
                logger.info("[generate-resume-preview] HIT PG sessionactivity — cid=%s", cid)

    if not cached:
        return {
            "success": False,
            "detail": "No resume data found for this session. Please upload a CV first.",
        }

    resume = build_structured_resume(cached, career_goal=career_goal, region=region)
    return {"success": True, "data": resume}
