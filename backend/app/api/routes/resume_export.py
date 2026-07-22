"""
Resume Export Route — Generates DOCX and PDF from Template.docx
Uses python-docx for DOCX and reportlab for PDF. No Office 365 dependency.
"""
from __future__ import annotations

import io
import os
import copy
from typing import Any

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib.colors import HexColor, black, white
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    ListFlowable, ListItem, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

router = APIRouter()

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "..", "..", "Template.docx")


# ===================== Request Schema =====================

class ContactInfo(BaseModel):
    name: str = ""
    headline: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    github: str = ""


class ExperienceItem(BaseModel):
    title: str = ""
    company: str = ""
    location: str = ""
    dates: str = ""
    bullets: list[str] = []


class EducationItem(BaseModel):
    degree: str = ""
    institution: str = ""
    dates: str = ""


class ResumeExportRequest(BaseModel):
    format: str = "docx"  # "docx" or "pdf"
    contact: ContactInfo = ContactInfo()
    summary: str = ""
    coreSkills: list[str] = []
    aiSkills: list[str] = []
    skillCategories: dict[str, list[str]] | None = None
    tools: list[str] = []
    competencies: list[str] = []
    certifications: list[str] = []
    experience: list[ExperienceItem] = []
    education: list[EducationItem] = []
    projects: list[str] = []


# ===================== DOCX Generation =====================

def _clear_cell(cell):
    """Remove all paragraphs from a cell except one empty one."""
    for i in range(len(cell.paragraphs) - 1, 0, -1):
        p = cell.paragraphs[i]._element
        p.getparent().remove(p)
    if cell.paragraphs:
        cell.paragraphs[0].clear()


def _add_paragraph(cell, text: str, bold: bool = False, size: int = 9, color: RGBColor | None = None, space_after: int = 2):
    """Add a paragraph to a cell with formatting."""
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def _add_heading_to_cell(cell, text: str):
    """Add a section heading (e.g. PROFESSIONAL SUMMARY)."""
    p = cell.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2D, 0x3A, 0x4A)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_bullet(cell, text: str):
    """Add a bullet point item."""
    p = cell.add_paragraph()
    run = p.add_run(f"•  {text}")
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Pt(8)


def _add_sidebar_heading(cell, text: str):
    """Add a sidebar category heading."""
    p = cell.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def _add_sidebar_item(cell, text: str):
    """Add a sidebar skill item."""
    p = cell.add_paragraph()
    run = p.add_run(f"›  {text}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)


def generate_docx(data: ResumeExportRequest) -> io.BytesIO:
    """Generate a DOCX resume from Template.docx filled with user data."""
    template_path = os.path.abspath(TEMPLATE_PATH)
    if not os.path.exists(template_path):
        raise HTTPException(status_code=500, detail="Template.docx not found")

    doc = Document(template_path)

    # ── Fill Table 0: Header ──
    header_table = doc.tables[0]

    # Row 0: Name
    _clear_cell(header_table.rows[0].cells[0])
    p = header_table.rows[0].cells[0].paragraphs[0]
    run = p.add_run(data.contact.name or "Your Name")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Row 1: Headline
    _clear_cell(header_table.rows[1].cells[0])
    p = header_table.rows[1].cells[0].paragraphs[0]
    run = p.add_run(data.contact.headline or "Professional Title")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Row 2: empty spacer (keep as-is)

    # Row 3: Contact info
    _clear_cell(header_table.rows[3].cells[0])
    contact_parts = []
    if data.contact.email:
        contact_parts.append(f"✉  {data.contact.email}")
    if data.contact.phone:
        contact_parts.append(f"✆  {data.contact.phone}")
    if data.contact.location:
        contact_parts.append(f"⊕  {data.contact.location}")
    if data.contact.linkedin:
        contact_parts.append(f"in  {data.contact.linkedin}")

    p = header_table.rows[3].cells[0].paragraphs[0]
    run = p.add_run("     ".join(contact_parts))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # ── Fill Table 1: Body ──
    body_table = doc.tables[1]
    sidebar_cell = body_table.rows[0].cells[0]
    main_cell = body_table.rows[0].cells[2]

    # Clear existing content
    _clear_cell(sidebar_cell)
    _clear_cell(main_cell)

    # ── SIDEBAR: Skill categories ──
    if data.skillCategories:
        for category, skills in data.skillCategories.items():
            _add_sidebar_heading(sidebar_cell, category)
            for skill in skills:
                _add_sidebar_item(sidebar_cell, skill)
    elif data.coreSkills:
        _add_sidebar_heading(sidebar_cell, "CORE SKILLS")
        for skill in data.coreSkills:
            _add_sidebar_item(sidebar_cell, skill)

    if data.tools:
        _add_sidebar_heading(sidebar_cell, "TOOLS")
        for tool in data.tools:
            _add_sidebar_item(sidebar_cell, tool)

    if data.competencies:
        _add_sidebar_heading(sidebar_cell, "COMPETENCIES")
        for comp in data.competencies:
            _add_sidebar_item(sidebar_cell, comp)

    if data.certifications:
        _add_sidebar_heading(sidebar_cell, "CERTIFICATIONS")
        for cert in data.certifications:
            _add_sidebar_item(sidebar_cell, cert)

    if data.aiSkills:
        _add_sidebar_heading(sidebar_cell, "AI SKILLS")
        for skill in data.aiSkills:
            _add_sidebar_item(sidebar_cell, skill)

    # ── MAIN CONTENT ──
    # Professional Summary
    if data.summary:
        _add_heading_to_cell(main_cell, "PROFESSIONAL SUMMARY")
        _add_paragraph(main_cell, data.summary, size=9, space_after=6)

    # Professional Experience
    if data.experience:
        _add_heading_to_cell(main_cell, "PROFESSIONAL EXPERIENCE")
        for exp in data.experience:
            title_line = exp.title
            if exp.company:
                title_line += f" — {exp.company}"
            if exp.dates:
                title_line += f"   {exp.dates}"
            _add_paragraph(main_cell, title_line, bold=True, size=9, space_after=2)
            if exp.location:
                _add_paragraph(main_cell, exp.location, size=8, color=RGBColor(0x66, 0x66, 0x66), space_after=2)
            for bullet in exp.bullets:
                _add_bullet(main_cell, bullet)

    # Education
    if data.education:
        _add_heading_to_cell(main_cell, "EDUCATION")
        for edu in data.education:
            edu_line = edu.degree
            if edu.institution:
                edu_line += f" — {edu.institution}"
            if edu.dates:
                edu_line += f"   {edu.dates}"
            _add_paragraph(main_cell, edu_line, bold=False, size=9, space_after=3)

    # Projects
    if data.projects:
        _add_heading_to_cell(main_cell, "PROJECTS")
        for proj in data.projects:
            _add_bullet(main_cell, proj)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ===================== PDF Generation =====================

def generate_pdf(data: ResumeExportRequest) -> io.BytesIO:
    """Generate a PDF resume using reportlab. No Office dependency."""
    buffer = io.BytesIO()
    page_width, page_height = A4
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.5 * cm,
        bottomMargin=1.5 * cm,
    )

    # ── Styles ──
    styles = getSampleStyleSheet()

    style_name = ParagraphStyle(
        'ResumeName', parent=styles['Normal'],
        fontSize=18, leading=22, textColor=HexColor('#1e293b'),
        fontName='Helvetica-Bold', spaceAfter=2
    )
    style_headline = ParagraphStyle(
        'ResumeHeadline', parent=styles['Normal'],
        fontSize=11, leading=14, textColor=HexColor('#475569'),
        fontName='Helvetica', spaceAfter=4
    )
    style_contact = ParagraphStyle(
        'ResumeContact', parent=styles['Normal'],
        fontSize=8, leading=11, textColor=HexColor('#64748b'),
        fontName='Helvetica', spaceAfter=10
    )
    style_section_heading = ParagraphStyle(
        'SectionHeading', parent=styles['Normal'],
        fontSize=10, leading=13, textColor=HexColor('#4F46E5'),
        fontName='Helvetica-Bold', spaceBefore=12, spaceAfter=4,
        borderWidth=0, borderPadding=0,
    )
    style_body = ParagraphStyle(
        'ResumeBody', parent=styles['Normal'],
        fontSize=9, leading=12, textColor=HexColor('#334155'),
        fontName='Helvetica', spaceAfter=3, alignment=TA_JUSTIFY
    )
    style_bold = ParagraphStyle(
        'ResumeBold', parent=styles['Normal'],
        fontSize=9, leading=12, textColor=HexColor('#1e293b'),
        fontName='Helvetica-Bold', spaceAfter=2
    )
    style_bullet = ParagraphStyle(
        'ResumeBullet', parent=styles['Normal'],
        fontSize=9, leading=12, textColor=HexColor('#334155'),
        fontName='Helvetica', leftIndent=12, spaceAfter=2
    )
    style_sidebar_heading = ParagraphStyle(
        'SidebarHeading', parent=styles['Normal'],
        fontSize=9, leading=11, textColor=HexColor('#4F46E5'),
        fontName='Helvetica-Bold', spaceBefore=8, spaceAfter=2
    )
    style_sidebar_item = ParagraphStyle(
        'SidebarItem', parent=styles['Normal'],
        fontSize=8, leading=10, textColor=HexColor('#475569'),
        fontName='Helvetica', spaceAfter=1
    )

    # ── Build content ──
    elements = []

    # Header
    elements.append(Paragraph(data.contact.name or "Your Name", style_name))
    elements.append(Paragraph(data.contact.headline or "", style_headline))

    contact_parts = []
    if data.contact.email:
        contact_parts.append(data.contact.email)
    if data.contact.phone:
        contact_parts.append(data.contact.phone)
    if data.contact.location:
        contact_parts.append(data.contact.location)
    if data.contact.linkedin:
        contact_parts.append(data.contact.linkedin)
    if contact_parts:
        elements.append(Paragraph("  |  ".join(contact_parts), style_contact))

    elements.append(Spacer(1, 6))

    # Build two-column layout: sidebar | main
    sidebar_elements = []
    main_elements = []

    # ── Sidebar content ──
    if data.skillCategories:
        for category, skills in data.skillCategories.items():
            sidebar_elements.append(Paragraph(category.upper(), style_sidebar_heading))
            for skill in skills:
                sidebar_elements.append(Paragraph(f"›  {skill}", style_sidebar_item))
    elif data.coreSkills:
        sidebar_elements.append(Paragraph("CORE SKILLS", style_sidebar_heading))
        for skill in data.coreSkills:
            sidebar_elements.append(Paragraph(f"›  {skill}", style_sidebar_item))

    if data.tools:
        sidebar_elements.append(Paragraph("TOOLS", style_sidebar_heading))
        for tool in data.tools:
            sidebar_elements.append(Paragraph(f"›  {tool}", style_sidebar_item))

    if data.competencies:
        sidebar_elements.append(Paragraph("COMPETENCIES", style_sidebar_heading))
        for comp in data.competencies:
            sidebar_elements.append(Paragraph(f"›  {comp}", style_sidebar_item))

    if data.certifications:
        sidebar_elements.append(Paragraph("CERTIFICATIONS", style_sidebar_heading))
        for cert in data.certifications:
            sidebar_elements.append(Paragraph(f"›  {cert}", style_sidebar_item))

    if data.aiSkills:
        sidebar_elements.append(Paragraph("AI SKILLS", style_sidebar_heading))
        for skill in data.aiSkills:
            sidebar_elements.append(Paragraph(f"›  {skill}", style_sidebar_item))

    # ── Main content ──
    if data.summary:
        main_elements.append(Paragraph("PROFESSIONAL SUMMARY", style_section_heading))
        main_elements.append(Paragraph(data.summary, style_body))

    if data.experience:
        main_elements.append(Paragraph("PROFESSIONAL EXPERIENCE", style_section_heading))
        for exp in data.experience:
            title_line = f"<b>{exp.title}</b>"
            if exp.company:
                title_line += f" — {exp.company}"
            if exp.dates:
                title_line += f"  <i>({exp.dates})</i>"
            main_elements.append(Paragraph(title_line, style_body))
            if exp.location:
                main_elements.append(Paragraph(f"<i>{exp.location}</i>", style_sidebar_item))
            for bullet in exp.bullets:
                main_elements.append(Paragraph(f"•  {bullet}", style_bullet))
            main_elements.append(Spacer(1, 4))

    if data.education:
        main_elements.append(Paragraph("EDUCATION", style_section_heading))
        for edu in data.education:
            edu_line = f"<b>{edu.degree}</b>"
            if edu.institution:
                edu_line += f" — {edu.institution}"
            if edu.dates:
                edu_line += f"  <i>({edu.dates})</i>"
            main_elements.append(Paragraph(edu_line, style_body))

    if data.projects:
        main_elements.append(Paragraph("PROJECTS", style_section_heading))
        for proj in data.projects:
            main_elements.append(Paragraph(f"•  {proj}", style_bullet))

    # ── Compose two-column table ──
    available_width = page_width - 3 * cm
    sidebar_width = available_width * 0.30
    main_width = available_width * 0.68

    # Combine sidebar and main into table cells
    from reportlab.platypus import TableStyle as TS

    col_data = [[sidebar_elements, main_elements]]
    layout_table = Table(col_data, colWidths=[sidebar_width, main_width])
    layout_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (0, 0), 0),
        ('RIGHTPADDING', (0, 0), (0, 0), 8),
        ('LEFTPADDING', (1, 0), (1, 0), 8),
        ('RIGHTPADDING', (1, 0), (1, 0), 0),
        ('LINEAFTER', (0, 0), (0, -1), 0.5, HexColor('#E2E8F0')),
    ]))

    elements.append(layout_table)

    doc.build(elements)
    buffer.seek(0)
    return buffer


# ===================== API Endpoint =====================

@router.post("/resume-export")
async def export_resume(request: ResumeExportRequest):
    """Generate and return a resume file (DOCX or PDF) from template."""
    fmt = request.format.lower().strip()

    if fmt == "docx":
        buffer = generate_docx(request)
        filename = f"{(request.contact.name or 'Resume').replace(' ', '_')}_Resume.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    elif fmt == "pdf":
        buffer = generate_pdf(request)
        filename = f"{(request.contact.name or 'Resume').replace(' ', '_')}_Resume.pdf"
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}. Use 'docx' or 'pdf'.")
