"""
Stage 1 - File Upload, Format Detection and Text Extraction

Input:  Uploaded file (PDF / DOCX / TXT), max 10 MB.
Output: raw_text (single flat string), request_id, content_hash, file_format.

Format detection: filename extension + MIME type + magic bytes (first 8 bytes).
PDF extractor:    PyMuPDF (fitz) primary, pdfplumber fallback.
DOCX extractor:   python-docx (paragraphs + table cells + headers/footers).
TXT extractor:    Encoding cascade UTF-8, UTF-16, latin-1. BOM stripped.
Cache:            Redis GET match_result:{sha256} on hit returns cached JSON.
"""
from __future__ import annotations

import hashlib
import io
import logging
import time
from typing import Any
from uuid import uuid4

from app.services.certification_pipeline.schemas import Stage1Output

logger = logging.getLogger(__name__)

# -- Constants --
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB

_PDF_MAGIC = b"%PDF"
_DOCX_MAGIC = b"PK\x03\x04"  # ZIP (DOCX is a ZIP archive)

_MIME_PDF = {"application/pdf"}
_MIME_DOCX = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/zip",
}
_MIME_TXT = {"text/plain", "application/octet-stream"}


# -- Format Detection --
def detect_format(filename: str, content_type: str | None, magic: bytes) -> str:
    """
    Detect file format from filename extension + MIME type + magic bytes.
    Returns one of: 'pdf', 'docx', 'txt'.
    Raises ValueError for unsupported formats.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # Magic bytes take priority
    if magic[:4] == _PDF_MAGIC:
        return "pdf"
    if magic[:4] == _DOCX_MAGIC and ext == "docx":
        return "docx"

    # Extension-based detection
    if ext == "pdf":
        return "pdf"
    if ext == "docx":
        return "docx"
    if ext in ("txt", "text"):
        return "txt"

    # MIME type fallback
    ct = (content_type or "").lower()
    if ct in _MIME_PDF:
        return "pdf"
    if ct in _MIME_DOCX:
        return "docx"
    if ct in _MIME_TXT:
        return "txt"

    raise ValueError(
        f"Unsupported file format: extension='{ext}', "
        f"content_type='{content_type}', magic={magic[:8].hex()}"
    )


# -- PDF Extraction --
def _extract_pdf_pymupdf(data: bytes) -> str:
    """Primary PDF extractor using PyMuPDF (fitz) - C-native, reading-order aware."""
    import fitz  # PyMuPDF

    pages: list[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text = page.get_text("text")
            if text and text.strip():
                pages.append(text)
    return "\n".join(pages)


def _extract_pdf_pdfplumber(data: bytes) -> str:
    """Fallback PDF extractor for complex table-embedded text."""
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text)
    return "\n".join(pages)


def extract_pdf(data: bytes) -> tuple[str, str]:
    """
    Extract text from PDF. Returns (text, extractor_used).
    Tries PyMuPDF first, falls back to pdfplumber if output is garbled/empty.
    """
    try:
        text = _extract_pdf_pymupdf(data)
        if text and len(text.strip()) > 20:
            return text, "pymupdf"
        logger.info("PyMuPDF produced insufficient text (%d chars), trying pdfplumber", len(text))
    except Exception as exc:
        logger.warning("PyMuPDF failed: %s - trying pdfplumber", exc)

    text = _extract_pdf_pdfplumber(data)
    if not text or not text.strip():
        raise ValueError("PDF text extraction failed - both PyMuPDF and pdfplumber produced empty output")
    return text, "pdfplumber"


# -- DOCX Extraction --
def extract_docx(data: bytes) -> tuple[str, str]:
    """
    Extract text from DOCX - paragraphs + table cells + headers/footers.
    Returns (text, extractor_used).
    """
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Table cells (cert names in two-column resume layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    # Headers and footers
    for section in doc.sections:
        for header in (section.header, section.first_page_header):
            if header:
                for para in header.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)
        for footer in (section.footer, section.first_page_footer):
            if footer:
                for para in footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

    raw_text = "\n".join(parts)
    if not raw_text.strip():
        raise ValueError("DOCX text extraction produced empty output")
    return raw_text, "python-docx"


# -- TXT Extraction --
def extract_txt(data: bytes) -> tuple[str, str]:
    """
    Decode text file with encoding cascade: UTF-8, UTF-16, latin-1.
    BOM is stripped. Returns (text, extractor_used).
    """
    # Strip BOM if present
    if data[:3] == b"\xef\xbb\xbf":
        data = data[3:]
    elif data[:2] in (b"\xff\xfe", b"\xfe\xff"):
        data = data[2:]

    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            text = data.decode(encoding)
            if text.strip():
                return text, f"txt-{encoding}"
        except (UnicodeDecodeError, UnicodeError):
            continue

    raise ValueError("TXT decoding failed - tried UTF-8, UTF-16, latin-1")


# -- Cache Helpers --
async def check_redis_cache(redis_client: Any, content_hash: str) -> dict[str, Any] | None:
    """Check Redis for a cached match result. Returns parsed JSON or None."""
    if redis_client is None:
        return None
    try:
        import orjson
        cached = await redis_client.get(f"match_result:{content_hash}")
        if cached:
            logger.info("Cache HIT for hash %s", content_hash[:16])
            return orjson.loads(cached)
    except Exception as exc:
        logger.warning("Redis cache check failed: %s", exc)
    return None


# -- Main Stage 1 Entry Point --
async def stage1_file_ingest(
    file_data: bytes,
    filename: str,
    content_type: str | None,
    redis_client: Any = None,
) -> Stage1Output:
    """
    Stage 1: File Upload, Format Detection and Text Extraction.

    1. Validate file size (max 10 MB).
    2. Detect file format via extension + MIME + magic bytes.
    3. Generate request_id (UUID4) and content_hash (SHA-256).
    4. Check Redis cache for existing result.
    5. Route to format-specific text extractor.
    6. Return Stage1Output with raw_text and metadata.

    Raises ValueError for unsupported formats, oversized files, or empty extractions.
    """
    t0 = time.perf_counter()

    # 1. Validate file size
    file_size = len(file_data)
    if file_size > MAX_FILE_SIZE:
        raise ValueError(f"File too large: {file_size:,} bytes (max {MAX_FILE_SIZE:,})")
    if file_size == 0:
        raise ValueError("Empty file uploaded")

    # 2. Detect format
    magic = file_data[:8]
    file_format = detect_format(filename, content_type, magic)

    # 3. Generate identifiers
    request_id = str(uuid4())
    content_hash = hashlib.sha256(file_data).hexdigest()

    # 4. Check Redis cache
    cached = await check_redis_cache(redis_client, content_hash)
    if cached is not None:
        elapsed_ms = (time.perf_counter() - t0) * 1000
        return Stage1Output(
            raw_text="",  # not needed for cache hit
            file_format=file_format,
            file_size_bytes=file_size,
            request_id=request_id,
            content_hash=content_hash,
            stageoutput={
                "cache_hit": True,
                "cached_response": cached,
                "file_format": file_format,
                "file_size_bytes": file_size,
                "content_hash": content_hash,
                "execution_ms": round(elapsed_ms, 2),
            },
        )

    # 5. Route to extractor
    extractor_used: str
    extraction_warning: str | None = None

    if file_format == "pdf":
        raw_text, extractor_used = extract_pdf(file_data)
    elif file_format == "docx":
        raw_text, extractor_used = extract_docx(file_data)
    elif file_format == "txt":
        raw_text, extractor_used = extract_txt(file_data)
    else:
        raise ValueError(f"No extractor for format: {file_format}")

    char_count = len(raw_text)
    if char_count < 10:
        raise ValueError(f"Extracted text too short ({char_count} chars) - file may be image-only or corrupt")

    # Check for potential PDF image-only (very low text density)
    if file_format == "pdf" and char_count < 50:
        extraction_warning = "Very low text density - PDF may be scanned/image-based"

    elapsed_ms = (time.perf_counter() - t0) * 1000

    # 6. Build stage output
    stageoutput: dict[str, Any] = {
        "file_format": file_format,
        "file_size_bytes": file_size,
        "char_count": char_count,
        "page_count": None,
        "extractor_used": extractor_used,
        "extraction_warning": extraction_warning,
        "content_hash": content_hash,
        "cache_hit": False,
        "execution_ms": round(elapsed_ms, 2),
    }

    # Get page count for PDF
    if file_format == "pdf":
        try:
            import fitz
            with fitz.open(stream=file_data, filetype="pdf") as doc:
                stageoutput["page_count"] = len(doc)
        except Exception:
            pass

    logger.info(
        "Stage 1 complete: format=%s, extractor=%s, chars=%d, %.1fms",
        file_format, extractor_used, char_count, elapsed_ms,
    )

    return Stage1Output(
        raw_text=raw_text,
        file_format=file_format,
        file_size_bytes=file_size,
        request_id=request_id,
        content_hash=content_hash,
        stageoutput=stageoutput,
    )
