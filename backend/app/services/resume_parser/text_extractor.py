"""Extract plain text from PDF, DOCX, DOC, ODS, and TXT resume files."""
from __future__ import annotations

import re
from io import BytesIO


# ── Spaced-text normalizer ───────────────────────────────────────────────────
# Stylised PDFs often render headings / names with a space between every
# character, e.g. "N S I K A K  E T U K".  Double-spaces denote word
# boundaries.  This collapses them back to normal text.

_SPACED_CHARS_RE = re.compile(r'^([A-Za-z0-9&/] ){2,}[A-Za-z0-9&/]$')


def _normalize_spaced_text(text: str) -> str:
    """Collapse spaced-out single characters in stylised PDF resumes.

    Examples
    --------
    >>> _normalize_spaced_text("N S I K A K  E T U K")
    'NSIKAK ETUK'
    >>> _normalize_spaced_text("P R O F E S S I O N A L  E X P E R I E N C E")
    'PROFESSIONAL EXPERIENCE'
    >>> _normalize_spaced_text("0 7 7 7 8 6 9 1 8 4 4")
    '07778691844'
    """
    lines = text.split('\n')
    result: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            result.append('')
            continue

        # Split on 2+ spaces (word boundaries in spaced text)
        segments = re.split(r' {2,}', stripped)
        collapsed: list[str] = []
        all_spaced = True

        for seg in segments:
            seg = seg.strip()
            if not seg:
                continue
            if _SPACED_CHARS_RE.match(seg):
                collapsed.append(seg.replace(' ', ''))
            else:
                all_spaced = False
                break

        if all_spaced and collapsed:
            result.append(' '.join(collapsed))
        else:
            result.append(stripped)

    return '\n'.join(result)


def extract_text(data: bytes, filename: str) -> str:
    """Return plain text from the uploaded resume bytes.

    Raises ValueError with a user-facing message on extraction failure.
    """
    name = filename.lower()

    if name.endswith(".txt"):
        return _from_txt(data)
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".docx"):
        return _from_docx(data)
    if name.endswith(".doc"):
        return _from_doc(data)
    if name.endswith(".ods"):
        return _from_ods(data)

    raise ValueError(f"Unsupported file extension for: {filename}")


# ── Format helpers ────────────────────────────────────────────────────────────

def _from_txt(data: bytes) -> str:
    try:
        import ftfy
        raw = data.decode("utf-8", errors="replace")
        return ftfy.fix_text(raw)
    except ImportError:
        return data.decode("utf-8", errors="replace")


def _find_column_split(text_blocks: list, page_width: float) -> float | None:
    """Find the column boundary by detecting the largest x-position gap.

    Returns the midpoint of the largest gap if it suggests a genuine
    multi-column layout, otherwise None (single column).
    """
    if len(text_blocks) < 6:
        return None

    x_starts = sorted(set(round(b[0]) for b in text_blocks))
    if len(x_starts) < 2:
        return None

    best_gap = 0
    best_mid = None
    for i in range(1, len(x_starts)):
        gap = x_starts[i] - x_starts[i - 1]
        if gap > best_gap:
            best_gap = gap
            best_mid = (x_starts[i - 1] + x_starts[i]) / 2

    # The gap must be significant (> 8% of page width) and the split
    # point must not be at the extreme edges (between 10%-80% of width).
    min_gap = page_width * 0.08
    if best_gap >= min_gap and best_mid is not None:
        if page_width * 0.10 < best_mid < page_width * 0.80:
            return best_mid

    return None


def _fitz_column_aware(data: bytes) -> str:
    """Extract text from PDF using PyMuPDF with column-aware block sorting.

    For multi-column resumes (e.g. sidebar + main), naive top-to-bottom
    extraction interleaves columns.  This groups text blocks into columns
    based on their x-position and reads each column fully before the next.

    Column boundary is detected dynamically via the largest gap in block
    x-positions rather than a fixed midpoint heuristic.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(stream=data, filetype="pdf")
    all_text: list[str] = []

    for page in doc:
        blocks = page.get_text("blocks")  # (x0, y0, x1, y1, text, block_no, block_type)
        text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]
        if not text_blocks:
            continue

        page_width = page.rect.width
        split_x = _find_column_split(text_blocks, page_width)

        if split_x is not None:
            left = [b for b in text_blocks if b[0] < split_x]
            right = [b for b in text_blocks if b[0] >= split_x]
            if left and right and len(left) >= 3 and len(right) >= 3:
                # Multi-column: read left column top-to-bottom, then right
                left.sort(key=lambda b: (b[1], b[0]))
                right.sort(key=lambda b: (b[1], b[0]))
                left_text = "\n".join(b[4].strip() for b in left)
                right_text = "\n".join(b[4].strip() for b in right)
                page_text = left_text + "\n\n" + right_text
            else:
                text_blocks.sort(key=lambda b: (b[1], b[0]))
                page_text = "\n".join(b[4].strip() for b in text_blocks)
        else:
            # Single column: standard top-to-bottom
            text_blocks.sort(key=lambda b: (b[1], b[0]))
            page_text = "\n".join(b[4].strip() for b in text_blocks)

        all_text.append(page_text)

    doc.close()
    return "\n".join(all_text).strip()


def _from_pdfplumber(data: bytes) -> str:
    """Extract text using pdfplumber — excellent for structured / styled PDFs."""
    import pdfplumber

    with pdfplumber.open(BytesIO(data)) as pdf:
        pages = [p.extract_text() or "" for p in pdf.pages]
    return "\n".join(pages).strip()


def _from_pdf(data: bytes) -> str:
    # Try PyMuPDF column-aware extraction first (handles multi-column resumes)
    try:
        text = _fitz_column_aware(data)
        if text and len(text.strip()) > 50:
            return _normalize_spaced_text(text)
    except Exception:
        pass

    # pdfplumber — good at preserving reading order in styled PDFs
    try:
        text = _from_pdfplumber(data)
        if text and len(text.strip()) > 50:
            return _normalize_spaced_text(text)
    except Exception:
        pass

    # Fall back to standard PyMuPDF
    try:
        import fitz
        doc = fitz.open(stream=data, filetype="pdf")
        pages = [page.get_text("text") for page in doc]
        doc.close()
        text = "\n".join(pages).strip()
        if text:
            return _normalize_spaced_text(text)
    except Exception:
        pass

    try:
        import pypdf
        reader = pypdf.PdfReader(BytesIO(data))
        pages = [p.extract_text() or "" for p in reader.pages]
        text = "\n".join(pages).strip()
        if text:
            return _normalize_spaced_text(text)
    except Exception:
        pass

    try:
        from pdfminer.high_level import extract_text as pm_extract
        text = pm_extract(BytesIO(data)).strip()
        if text:
            return _normalize_spaced_text(text)
    except Exception:
        pass

    raise ValueError("Could not extract text from the PDF. The file may be scanned or password-protected.")


def _from_docx(data: bytes) -> str:
    try:
        import docx
        document = docx.Document(BytesIO(data))
        lines = [p.text for p in document.paragraphs if p.text.strip()]
        # Also grab text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)
    except ImportError:
        pass

    try:
        import docx2txt
        return docx2txt.process(BytesIO(data))
    except ImportError:
        pass

    raise ValueError("Could not extract text from the DOCX file. Ensure python-docx is installed.")


def _from_doc(data: bytes) -> str:
    # .doc (legacy Word) — try docx2txt, then antiword via subprocess
    try:
        import docx2txt
        text = docx2txt.process(BytesIO(data))
        if text and text.strip():
            return text
    except Exception:
        pass

    try:
        import subprocess
        import tempfile, os
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        result = subprocess.run(
            ["antiword", tmp_path],
            capture_output=True, text=True, timeout=15
        )
        os.unlink(tmp_path)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except Exception:
        pass

    raise ValueError(
        "Could not extract text from the .doc file. "
        "Please convert it to DOCX or PDF and try again."
    )


def _from_ods(data: bytes) -> str:
    try:
        from odf.opendocument import load
        from odf import table as odf_table, text as odf_text
        doc = load(BytesIO(data))
        lines: list[str] = []
        for tbl in doc.getElementsByType(odf_table.Table):
            for row in tbl.getElementsByType(odf_table.TableRow):
                row_cells: list[str] = []
                for cell in row.getElementsByType(odf_table.TableCell):
                    ps = cell.getElementsByType(odf_text.P)
                    cell_text = " ".join(
                        str(t.firstChild.data) if getattr(t, "firstChild", None) else ""
                        for t in ps
                    ).strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    lines.append("\t".join(row_cells))
        return "\n".join(lines)
    except ImportError:
        raise ValueError("ODS support requires 'odfpy'. Please install it.")
