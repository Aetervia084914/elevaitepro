"""Stage 2 — Document -> Raw Text Extraction (No Section Detection)."""
from __future__ import annotations

import io
import logging
import re
import subprocess
import tempfile
import time
from pathlib import Path

import ftfy
import regex
from langdetect import detect as langdetect_detect, LangDetectException

from app.schemas.skill_normalizer import FileFormat, PipelineContext, StageStatus

log = logging.getLogger(__name__)

MAX_TEXT_CHARS: int = 50_000


class UnsupportedLanguageError(Exception):
    """Raised when detected language is not English."""
    def __init__(self, detected: str):
        self.detected = detected
        super().__init__(f"Unsupported language: {detected!r}. Only English (en) is supported.")


class TextExtractionError(Exception):
    """Raised when text extraction fails for a supported format."""
    pass


# -- Format handlers ----------------------------------------------------------

def _extract_docx(file_bytes: bytes) -> tuple[str, str]:
    """Extract text from .docx using python-docx (paragraphs + table cells)."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts), "docx"


def _extract_doc(file_bytes: bytes) -> tuple[str, str]:
    """Extract text from legacy .doc."""
    try:
        import docx2txt
        text = docx2txt.process(io.BytesIO(file_bytes))
        if text and text.strip():
            return text.strip(), "doc_docx2txt"
    except Exception:
        pass

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            doc_path = Path(tmpdir) / "input.doc"
            doc_path.write_bytes(file_bytes)
            result = subprocess.run(
                [
                    "soffice", "--headless", "--convert-to", "txt:Text",
                    "--outdir", tmpdir, str(doc_path),
                ],
                capture_output=True,
                timeout=30,
            )
            txt_path = Path(tmpdir) / "input.txt"
            if txt_path.exists():
                text = txt_path.read_text(encoding="utf-8", errors="replace").strip()
                if text:
                    return text, "doc_libreoffice"
    except (FileNotFoundError, subprocess.TimeoutExpired) as exc:
        log.warning("LibreOffice conversion failed: %s", exc)

    raise TextExtractionError(
        "Could not extract text from .doc file. "
        "Install LibreOffice (soffice) for .doc support."
    )


def _extract_pdf(file_bytes: bytes) -> tuple[str, str]:
    """Extract text from PDF."""
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(io.BytesIO(file_bytes))
        if text and len(text.strip()) > 50:
            return text.strip(), "pdf_textlayer"
    except Exception as exc:
        log.debug("pdfminer extraction failed: %s", exc)

    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        parts: list[str] = []
        for page in doc:
            page_text = page.get_text()
            if page_text:
                parts.append(page_text)
        doc.close()
        text = "\n".join(parts).strip()
        if text and len(text) > 50:
            return text, "pdf_ocr"
    except Exception as exc:
        log.debug("PyMuPDF extraction failed: %s", exc)

    raise TextExtractionError(
        "Could not extract text from PDF. The file may be image-only "
        "without an OCR text layer."
    )


def _extract_odt(file_bytes: bytes) -> tuple[str, str]:
    """Extract text from .odt using odfpy."""
    from odf.opendocument import load as odf_load
    from odf.text import P as OdfP
    from odf import teletype

    doc = odf_load(io.BytesIO(file_bytes))
    parts: list[str] = []
    for para in doc.getElementsByType(OdfP):
        text = teletype.extractText(para).strip()
        if text:
            parts.append(text)

    return "\n".join(parts), "odt"


def _extract_ods(file_bytes: bytes) -> tuple[str, str]:
    """Extract text from .ods using odfpy (all cell text)."""
    from odf.opendocument import load as odf_load
    from odf.table import Table, TableRow, TableCell
    from odf import teletype

    doc = odf_load(io.BytesIO(file_bytes))
    parts: list[str] = []
    for table in doc.getElementsByType(Table):
        for row in table.getElementsByType(TableRow):
            for cell in row.getElementsByType(TableCell):
                text = teletype.extractText(cell).strip()
                if text:
                    parts.append(text)

    return "\n".join(parts), "ods"


def _extract_txt(file_bytes: bytes) -> tuple[str, str]:
    """Extract text from .txt — UTF-8 decode, BOM strip."""
    if file_bytes[:3] == b"\xef\xbb\xbf":
        file_bytes = file_bytes[3:]
    elif file_bytes[:2] in (b"\xff\xfe", b"\xfe\xff"):
        file_bytes = file_bytes[2:]

    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")

    return text.strip(), "txt"


# Dispatch table
_HANDLERS: dict[FileFormat, callable] = {
    FileFormat.DOCX: _extract_docx,
    FileFormat.DOC:  _extract_doc,
    FileFormat.PDF:  _extract_pdf,
    FileFormat.ODT:  _extract_odt,
    FileFormat.ODS:  _extract_ods,
    FileFormat.TXT:  _extract_txt,
}


# -- Text cleaning ------------------------------------------------------------

_CONTROL_CHAR_RE = regex.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")
_BULLET_RE = regex.compile(r"[\u2022\u25CF\u25CB\u25E6\u25AA\u25B8\u25BA\u25C6\u2605\u2726\u2727\u27A4\u27A2\u2192\u2023\u2043\u2010\u2013\u2014]")
_MULTI_SPACE_RE = re.compile(r"[^\S\n]+")
_MULTI_NEWLINE_RE = re.compile(r"\n{3,}")


def _clean_text(raw: str) -> str:
    """Clean extracted text."""
    text = ftfy.fix_text(raw)
    text = _CONTROL_CHAR_RE.sub("", text)
    text = _BULLET_RE.sub(" ", text)
    text = _MULTI_SPACE_RE.sub(" ", text)
    text = _MULTI_NEWLINE_RE.sub("\n\n", text)
    return text.strip()


# -- Language gate ------------------------------------------------------------

def _detect_language(text: str) -> str:
    """Detect language using langdetect."""
    sample = text[:2000]
    try:
        return langdetect_detect(sample)
    except LangDetectException:
        log.warning("Language detection failed — defaulting to 'en'.")
        return "en"


# -- Main entry point ---------------------------------------------------------

async def run_stage2(
    ctx: PipelineContext,
    file_bytes: bytes,
) -> PipelineContext:
    """Execute Stage 2 of the pipeline."""
    t0 = time.perf_counter()

    if ctx.file_format is None:
        raise TextExtractionError("file_format not set — Stage 1 must run first.")

    handler = _HANDLERS.get(ctx.file_format)
    if handler is None:
        raise TextExtractionError(f"No handler for format: {ctx.file_format.value}")

    # 2a. Extract raw text
    raw_text, method = handler(file_bytes)
    ctx.extraction_method = method

    if not raw_text or len(raw_text.strip()) < 10:
        raise TextExtractionError(
            "Extracted text is empty or too short (< 10 chars). "
            "The file may be image-only or corrupt."
        )

    # 2b. Clean text
    cleaned = _clean_text(raw_text)
    ctx.original_char_count = len(cleaned)

    # 2c. Text size cap
    if len(cleaned) > MAX_TEXT_CHARS:
        log.warning(
            "Stage 2: Text truncated from %d to %d chars (request_id=%s).",
            len(cleaned), MAX_TEXT_CHARS, ctx.request_id,
        )
        cleaned = cleaned[:MAX_TEXT_CHARS]
        ctx.truncated = True

    ctx.raw_text = cleaned

    # 2d. Language detection (EN gate)
    lang = _detect_language(cleaned)
    ctx.detected_language = lang

    duration_ms = int((time.perf_counter() - t0) * 1000)

    if lang != "en":
        ctx.add_stage_result(
            stage="2",
            status=StageStatus.REJECTED,
            duration_ms=duration_ms,
            payload={
                "char_count":          ctx.original_char_count,
                "detected_language":   lang,
                "extraction_method":   method,
            },
            error_message=f"UNSUPPORTED_LANGUAGE: {lang}",
        )
        raise UnsupportedLanguageError(lang)

    # 2e. Audit
    ctx.add_stage_result(
        stage="2",
        status=StageStatus.OK,
        duration_ms=duration_ms,
        payload={
            "char_count":          len(ctx.raw_text),
            "original_char_count": ctx.original_char_count,
            "truncated":           ctx.truncated,
            "detected_language":   lang,
            "extraction_method":   method,
        },
    )

    log.info(
        "Stage 2 OK — request_id=%s method=%s chars=%d lang=%s truncated=%s %dms",
        ctx.request_id, method, len(ctx.raw_text), lang, ctx.truncated, duration_ms,
    )
    return ctx
